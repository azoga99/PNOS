# -*- coding: utf-8 -*-
"""
Этап 2: Копирование таблиц.
Парсинг файла "Акт замеров толщины стенки.docx" из папки "Первичка" 
и перенос 3-й таблицы в "ПНОС_ТТП.xlsm" на листы "Импорт" и "Импорт (ПР)".
"""

import os
import time
import math
import difflib
import threading

from PySide6.QtCore import QThread, Signal

import pythoncom
import win32com.client
import win32gui
import win32con
import docx

# Импорт конфигурации и утилит из Этапа 3
from config import CONFIG
from workers.stage3_worker import DialogKiller, safe_close_com


def find_folder_fuzzy(base: str, target: str, cutoff: float = 0.7) -> str | None:
    """Нечёткий поиск подпапки."""
    if not os.path.exists(base):
        return None
    dirs = {d.lower(): d for d in os.listdir(base) if os.path.isdir(os.path.join(base, d))}
    
    if target.lower() in dirs:
        return os.path.join(base, dirs[target.lower()])
        
    m = difflib.get_close_matches(target.lower(), dirs.keys(), n=1, cutoff=cutoff)
    return os.path.join(base, dirs[m[0]]) if m else None

def find_file_fuzzy(folder: str, target: str, exts: tuple, cutoff: float = 0.7) -> str | None:
    """Нечёткий поиск файла по имени и расширению."""
    if not folder or not os.path.exists(folder):
        return None
    files = {}
    for f in os.listdir(folder):
        if f.startswith("~$"):
            continue
        p = os.path.join(folder, f)
        if os.path.isfile(p):
            name, ext = os.path.splitext(f)
            if ext.lower() in exts:
                files[name.lower()] = p
                
    if target.lower() in files:
        return files[target.lower()]
        
    m = difflib.get_close_matches(target.lower(), files.keys(), n=1, cutoff=cutoff)
    return files[m[0]] if m else None

def find_file_exact_ext(folder: str, exts: tuple) -> str | None:
    """Ищет первый попавшийся файл с заданным расширением."""
    if not folder or not os.path.exists(folder):
        return None
    for f in sorted(os.listdir(folder)):
        if f.startswith("~$"):
            continue
        p = os.path.join(folder, f)
        if os.path.isfile(p) and os.path.splitext(f)[1].lower() in exts:
            return p
    return None

def clean_word_text(text: str) -> str:
    """Очистка текста из ячеек таблицы."""
    if not text:
        return ""
    # python-docx отдает текст как есть, просто делаем strip от лишних пробелов по краям
    return text.strip()


class Stage2Worker(QThread):
    """Фоновый поток для Этапа 2 — копирование таблиц."""

    log = Signal(str)             # Сообщение в лог
    progress = Signal(int)        # Прогресс 0-100
    action_update = Signal(str)   # Текст для UI над прогресс-баром
    report_ready = Signal(list)   # Список словарей для модального окна-отчёта
    table_preview = Signal(list)  # Срез для превью таблицы 3 (list of lists)
    finished_ok = Signal(bool)    # Завершение (True = успех)
    info = Signal(str, str)       # Дружелюбный статус (сообщение, категория)

    def __init__(self, local_root: str, parent=None):
        super().__init__(parent)
        self.local_root = local_root
        self._is_cancelled = False
        
        # Настройки поиска "Акта"
        self.doc_name_target = "Акт замеров толщины стенки"
        self.folder_name_target = "Первичка"

    def cancel(self):
        self._is_cancelled = True

    def run(self):
        self.log.emit("═" * 40)
        self.log.emit("ЭТАП 2: Копирование таблиц из Первички")
        self.log.emit("═" * 40)

        if not os.path.isdir(self.local_root):
            self.log.emit("❌ Корневая папка не найдена.")
            self.finished_ok.emit(False)
            return

        # Ищем все папки пунктов "п.*"
        point_folders = []
        for d in os.listdir(self.local_root):
            dpath = os.path.join(self.local_root, d)
            if os.path.isdir(dpath) and d.startswith("п."):
                point_folders.append(dpath)

        if not point_folders:
            self.log.emit("⚠ Не найдено локальных папок пунктов (п.*).")
            self.finished_ok.emit(False)
            return

        self.log.emit(f"📋 Найдено пунктов для обработки: {len(point_folders)}")
        self.progress.emit(5)

        # 2. Инициализация COM и Диалог Киллера
        pythoncom.CoInitialize()
        dk = DialogKiller(log_callback=self.log.emit)
        
        xl_app = None

        try:
            dk.start()
            
            self.log.emit("\n⚙ Инициализация Excel в фоне...")
            xl_app = win32com.client.Dispatch("Excel.Application")
            xl_app.Visible = False
            xl_app.DisplayAlerts = False
            xl_app.AskToUpdateLinks = False
            xl_app.ScreenUpdating = False
            xl_app.EnableEvents = False

            self.info.emit("Excel готов к работе (Word читается напрямую)", "done")
            self.progress.emit(10)
            self.log.emit("✅ Excel COM готов. Начинаем извлечение (чтение Word через python-docx).")

            processed_count = 0
            errors_count = 0
            report_data = []
            
            for i, p_folder in enumerate(point_folders):
                if self._is_cancelled:
                    self.log.emit("\n❌ Выполнение отменено пользователем!")
                    break

                p_name = os.path.basename(p_folder)
                pt_num = p_name.replace("п.", "")
                self.log.emit(f"\n── {p_name} ──")
                self.info.emit(f"Обработка п.{pt_num}...", "wait")
                self.action_update.emit(f"Обрабатываем пункт {pt_num}...")
                
                point_report = {
                    "№ Пункта": pt_num,
                    "Акт найден": False,
                    "Таблица 3 найдена": False,
                    "Результат": "Ожидание"
                }
                
                work_wb = None
                doc = None

                try:
                    # Поиск файла ПНОС_ТТП.xlsm
                    work_path = find_file_exact_ext(p_folder, (".xls", ".xlsx", ".xlsm"))
                    if not work_path:
                        self.log.emit(f"   ✗ Ошибка: Excel-файл (ПНОС_ТТП.xlsm) не найден в папке")
                        errors_count += 1
                        continue

                    # Поиск папки "Первичка"
                    pervichka_dir = find_folder_fuzzy(p_folder, self.folder_name_target)
                    if not pervichka_dir:
                        self.log.emit(f"   ✗ Ошибка: Папка «Первичка» не найдена")
                        errors_count += 1
                        continue
                        
                    # Поиск "Акт замеров толщины стенки.docx"
                    akt_path = find_file_fuzzy(pervichka_dir, self.doc_name_target, (".doc", ".docx"))
                    if not akt_path:
                        self.log.emit(f"   ✗ Ошибка: Файл «Акт замеров...» не найден в Первичке")
                        point_report["Результат"] = "❌ Нет Акта"
                        errors_count += 1
                        report_data.append(point_report)
                        continue

                    point_report["Акт найден"] = True
                    self.log.emit(f"   [Word] Читаем (python-docx): {os.path.basename(akt_path)}")
                    
                    # 1. Читаем таблицу №3 из Word
                    try:
                        doc = docx.Document(akt_path)
                    except Exception as e:
                        self.log.emit(f"   ✗ Ошибка открытия Word-документа: {e}")
                        point_report["Результат"] = "❌ Ошибка чтения .docx"
                        errors_count += 1
                        report_data.append(point_report)
                        continue
                    
                    if len(doc.tables) < 3:
                        msg = f"Ожидалось минимум 3 таблицы, найдено {len(doc.tables)}"
                        self.log.emit(f"   ✗ Ошибка в Акте: {msg}")
                        point_report["Результат"] = "❌ Меньше 3 таблиц"
                        errors_count += 1
                        doc = None
                        report_data.append(point_report)
                        continue
                        
                    point_report["Таблица 3 найдена"] = True
                    table = doc.tables[2]
                    
                    self.log.emit(f"   [Word] Найдена таблица 3: {len(table.rows)} строк")
                    
                    # Извлечение данных в двумерный массив массива (list of lists)
                    table_data = []
                    
                    for row in table.rows:
                        current_row = []
                        for cell in row.cells:
                            text = clean_word_text(cell.text)
                            current_row.append(text)
                            
                        # python-docx automatically handles merged cell logic by duplicating 
                        # text in the invisible cells of the merged block. This yields a perfect 
                        # rectangular grid structure matching Excel array expectations.
                        table_data.append(current_row)
                        
                    doc = None

                    if not table_data:
                        self.log.emit("   ⚠ Таблица 3 оказалась пустой")
                        processed_count += 1
                        continue
                        
                    # Отправляем превью таблицы (первые 5 строк)
                    preview_data = table_data[:5]
                    self.table_preview.emit(preview_data)

                    # 2. Пишем в Excel
                    self.log.emit("   [Excel] Открытие шаблона ПНОС_ТТП.xlsm...")
                    work_wb = xl_app.Workbooks.Open(
                        work_path, UpdateLinks=0,
                        IgnoreReadOnlyRecommended=True, Notify=False)

                    # Пишем на оба листа: "Импорт" и "Импорт (ПР)"
                    target_sheets = ["Импорт", "Импорт (ПР)"]
                    
                    max_cols = max(len(row) for row in table_data) if table_data else 0
                    if max_cols == 0:
                        raise ValueError("Массив данных пуст, хотя строк > 0")

                    # Подготавливаем 2D массив, чтобы Excel мог принять его в Range
                    # Выравниваем строки по длине, заполняя пустыми строками
                    win32_data = []
                    for row in table_data:
                        win32_row = tuple([str(cell) for cell in row] + [""] * (max_cols - len(row)))
                        win32_data.append(win32_row)
                    
                    win32_data = tuple(win32_data)

                    rows_n = len(win32_data)
                    cols_n = len(win32_data[0]) if rows_n > 0 else 0
                    
                    # A2 - это R2C1. Конечная ячейка R(1+rows_n)C(cols_n)
                    start_cell_addr = "A2"
                    
                    for sheet_name in target_sheets:
                        try:
                            ws = work_wb.Sheets(sheet_name)
                            # Запись построчно для максимальной надежности COM (массивы могут сбоить)
                            for r_idx, row_data in enumerate(win32_data):
                                # +2 потому что A2 это вторая строка
                                cell_start = ws.Cells(2 + r_idx, 1)
                                cell_end = ws.Cells(2 + r_idx, cols_n)
                                ws.Range(cell_start, cell_end).Value = row_data
                            self.log.emit(f"   ✓ Данные ({rows_n}x{cols_n}) вставлены на лист «{sheet_name}»")
                        except Exception as e:
                            self.log.emit(f"   ⚠ Ошибка записи на лист «{sheet_name}»: {e}")

                    # Сохранение
                    work_wb.Save()
                    work_wb.Close()
                    work_wb = None

                    self.log.emit(f"   ✓ Успех. Таблица №3 перенесена.")
                    self.info.emit(f"п.{pt_num} — таблица перенесена", "done")
                    point_report["Результат"] = "✅ Успешно"
                    processed_count += 1

                except Exception as ex:
                    self.log.emit(f"   ❌ Ошибка при обработке {p_name}: {ex}")
                    point_report["Результат"] = f"❌ Ошибка"
                    errors_count += 1
                finally:
                    report_data.append(point_report)
                    safe_close_com(work_wb, save=False)
                    safe_close_com(doc, save=False)

                # Шаг прогресса
                pct = 10 + int((i + 1) / len(point_folders) * 85)
                self.progress.emit(min(pct, 95))

            # 4. Финиш
            self.progress.emit(100)
            self.info.emit("Этап 2 завершен!", "done")
            self.log.emit(f"\n{'═' * 40}")
            self.log.emit(f"✅ Этап 2 завершён! Обработано успешно: {processed_count}, Ошибок: {errors_count}")

            # Сигнализируем успех, если не было сплошных ошибок
            success = not self._is_cancelled and (processed_count > 0 or len(point_folders) == 0)
            self.report_ready.emit(report_data)
            self.action_update.emit("Завершено!")
            self.finished_ok.emit(success)

        except Exception as e:
            self.info.emit(f"Ошибка этапа 2: {str(e)[:50]}", "error")
            self.log.emit(f"\n❌ Критическая ошибка Этапа 2: {e}")
            import traceback
            self.log.emit(traceback.format_exc())
            self.finished_ok.emit(False)

        finally:
            dk.stop()
            self.log.emit("🧹 Закрытие COM объектов...")
            
            if xl_app:
                try:
                    xl_app.ScreenUpdating = True
                    xl_app.Quit()
                except Exception:
                    pass
            
            pythoncom.CoUninitialize()

