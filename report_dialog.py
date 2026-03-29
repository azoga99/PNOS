# -*- coding: utf-8 -*-
import os
import json
from PySide6.QtWidgets import (
    QDialog, QVBoxLayout, QHBoxLayout, QPushButton, 
    QTableWidget, QTableWidgetItem, QLabel, QHeaderView, QAbstractItemView, QMessageBox,
    QTextEdit
)
from PySide6.QtCore import Qt
from PySide6.QtGui import QColor, QFont
import docx
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from status_manager import get_status_file_path

# Цвета статусов
COLOR_SUCCESS = QColor("#43A047") # Зеленый
COLOR_ERROR = QColor("#E53935")   # Красный
COLOR_SKIP = QColor("#9E9E9E")    # Серый

class ReportDialog(QDialog):
    """
    Финальное модальное окно с отчетом о проделанной работе.
    """
    def __init__(self, local_root, parent=None):
        super().__init__(parent)
        self.local_root = local_root
        self.setWindowTitle("Итоговый отчет о генерации ПНОС")
        self.resize(1000, 600)
        
        # Настройка стиля (Лаконичный и современный)
        self.setStyleSheet("""
            QDialog {
                background-color: #FAFAFA;
            }
            QLabel#Title {
                font-size: 22px;
                font-family: "Segoe UI", sans-serif;
                font-weight: bold;
                color: #2C3E50;
                padding: 10px 0;
            }
            QLabel#Summary {
                font-size: 14px;
                font-family: "Segoe UI", sans-serif;
                color: #34495E;
                padding-bottom: 20px;
            }
            QTableWidget {
                background-color: white;
                border: 1px solid #E0E0E0;
                border-radius: 6px;
                gridline-color: transparent; /* Убираем стандартную сетку */
                outline: none;
            }
            QHeaderView::section {
                background-color: #F8F9FA;
                border: none;
                border-bottom: 2px solid #E0E0E0;
                padding: 12px;
                font-weight: bold;
                font-family: "Segoe UI", sans-serif;
                font-size: 13px;
                color: #555555;
            }
            QTableWidget::item {
                padding: 10px;
                border-bottom: 1px solid #F0F0F0;
                font-family: "Segoe UI", sans-serif;
            }
            QPushButton {
                background-color: #2980B9;
                color: white;
                border: none;
                border-radius: 5px;
                padding: 12px 24px;
                font-family: "Segoe UI", sans-serif;
                font-size: 14px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #3498DB;
            }
            QPushButton#CloseBtn {
                background-color: #95A5A6;
            }
            QPushButton#CloseBtn:hover {
                background-color: #7F8C8D;
            }
        """)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(30, 30, 30, 30)
        
        # Заголовок
        title = QLabel("Итоговый отчёт ПНОС")
        title.setObjectName("Title")
        title.setAlignment(Qt.AlignCenter)
        layout.addWidget(title)
        
        self.summary_lbl = QLabel()
        self.summary_lbl.setObjectName("Summary")
        self.summary_lbl.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.summary_lbl)
        
        # Таблица
        self.table = QTableWidget()
        self.table.setColumnCount(6)
        self.table.setHorizontalHeaderLabels([
            "№ Пункта", "Этап 1 (Загрузка)", "Этап 2 (Извлечение)", 
            "Этап 3 (Макросы)", "Этап 4 (Слияние)", "Этап 5 (Фотография)"
        ])
        
        # Настройки взаимодействия
        self.table.setSelectionMode(QAbstractItemView.NoSelection)
        self.table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.table.verticalHeader().setVisible(False)
        self.table.setShowGrid(False) # Полностью убираем дефолтную сетку, оставляя border-bottom
        
        header = self.table.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.ResizeToContents)
        for i in range(1, 6):
            header.setSectionResizeMode(i, QHeaderView.Stretch)
            
        layout.addWidget(self.table)

        # Поле для предупреждений / отклонений (скрыто, если их нет)
        self.warnings_text = QTextEdit()
        self.warnings_text.setReadOnly(True)
        self.warnings_text.setStyleSheet("background-color: #FFF3F3; color: #D32F2F; font-size: 13px; font-weight: bold; border: 1px solid #FFCDD2; border-radius: 6px;")
        self.warnings_text.setMaximumHeight(80)
        self.warnings_text.hide()
        layout.addWidget(self.warnings_text)
        
        # Кнопки
        btn_layout = QHBoxLayout()
        btn_layout.addStretch()
        
        self.btn_export = QPushButton("Выгрузить в DOCX")
        self.btn_export.clicked.connect(self.export_to_docx)
        btn_layout.addWidget(self.btn_export)
        
        self.btn_close = QPushButton("Закрыть")
        self.btn_close.setObjectName("CloseBtn")
        self.btn_close.clicked.connect(self.accept)
        btn_layout.addWidget(self.btn_close)
        
        btn_layout.addStretch()
        layout.addLayout(btn_layout)
        
        # Загрузка
        self.load_data()

    def get_symbol_and_color(self, status):
        """Возвращает строку с кружком и цвет для QTableWidgetItem."""
        if status is True:
            return "🟢 Готово", COLOR_SUCCESS
        elif status is False:
            return "🔴 Отсутствует", COLOR_ERROR
        return "⚪ Нет данных", COLOR_SKIP

    def load_data(self):
        if not os.path.exists(self.local_root):
            self.summary_lbl.setText("Папка проекта не найдена.")
            return

        point_folders = [d for d in os.listdir(self.local_root) 
                         if os.path.isdir(os.path.join(self.local_root, d)) and d.startswith("п.")]
        
        # Сортировка (п.1, п.2, п.10)
        def sort_key(s):
            try: return int(s.replace("п.", ""))
            except: return 999999
        point_folders.sort(key=sort_key)
        
        self.table.setRowCount(len(point_folders))
        
        completed_all = 0
        self.warnings_list = []
        
        bold_font = QFont()
        bold_font.setBold(True)
        
        for row_idx, folder_name in enumerate(point_folders):
            pt_num = folder_name.replace("п.", "")
            folder_path = os.path.join(self.local_root, folder_name)
            status_path = get_status_file_path(folder_path)
            
            # Чтение статуса
            st1, st2, st3, st4, st5 = False, False, False, False, False
            if os.path.exists(status_path):
                try:
                    with open(status_path, "r", encoding="utf-8") as f:
                        data = json.load(f)
                        st1 = data.get("stage1", False)
                        st2 = data.get("stage2", False)
                        st3 = data.get("stage3", False)
                        st4 = data.get("stage4", False)
                        st5 = data.get("stage5", False)
                        
                        warning = data.get("stage5_warning", "")
                        if warning:
                            self.warnings_list.append(warning)
                except: pass

            if st1 and st2 and st3 and st4 and st5:
                completed_all += 1
                
            # Заполняем строку
            # Столбец 0: Номер
            it_num = QTableWidgetItem(pt_num)
            it_num.setTextAlignment(Qt.AlignCenter)
            it_num.setFont(bold_font)
            self.table.setItem(row_idx, 0, it_num)
            
            # Столбцы статусов
            statuses = [st1, st2, st3, st4, st5]
            for col_idx, s_val in enumerate(statuses, start=1):
                text, color = self.get_symbol_and_color(s_val)
                it = QTableWidgetItem(text)
                it.setTextAlignment(Qt.AlignCenter)
                it.setForeground(color)
                it.setFont(bold_font)
                self.table.setItem(row_idx, col_idx, it)
                
        self.summary_lbl.setText(f"Всего пунктов: {len(point_folders)}   |   Полностью готовы: <span style='color:#43A047; font-weight:bold'>{completed_all}</span>   |   Ожидают внимания: <span style='color:#E53935; font-weight:bold'>{len(point_folders) - completed_all}</span>")
        self.summary_lbl.setTextFormat(Qt.RichText)
        
        if self.warnings_list:
            self.warnings_text.setText("\n".join(self.warnings_list))
            self.warnings_text.show()
        else:
            self.warnings_text.hide()

    def export_to_docx(self):
        """Экспорт таблицы отчета в файл .docx."""
        try:
            doc = docx.Document()
            
            # Изменение ориентации страницы на альбомную для широкой таблицы
            section = doc.sections[-1]
            new_width, new_height = section.page_height, section.page_width
            section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
            section.page_width = new_width
            section.page_height = new_height
            
            # Заголовок
            h1 = doc.add_paragraph("Отчёт о генерации пакета ПНОС")
            h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in h1.runs:
                run.font.bold = True
                run.font.size = Pt(16)
                
            # Сводка (парсинг RichText вызовет сложности у docx, поэтому пишем сырой текст)
            import re
            raw_summary = re.sub(r'<[^>]+>', '', self.summary_lbl.text())
            summary = doc.add_paragraph(raw_summary)
            summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Создаем таблицу
            rows = self.table.rowCount()
            cols = self.table.columnCount()
            word_table = doc.add_table(rows=rows + 1, cols=cols)
            word_table.style = 'Table Grid'
            
            # Шапка
            for c in range(cols):
                cell = word_table.cell(0, c)
                cell.text = self.table.horizontalHeaderItem(c).text()
                for rp in cell.paragraphs:
                    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in rp.runs:
                        run.font.bold = True
                        
            # Тело
            for r in range(rows):
                for c in range(cols):
                    item = self.table.item(r, c)
                    text = item.text() if item else ""
                    # Убираем цветные кружочки из экспорта для чистоты, оставляя только текст
                    text = text.replace("🟢 ", "").replace("🔴 ", "").replace("⚪ ", "")
                    
                    cell = word_table.cell(r + 1, c)
                    cell.text = text
                    for rp in cell.paragraphs:
                        rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
            # Добавление предупреждений (отклонений) под таблицей
            if self.warnings_list:
                doc.add_paragraph("") # пустая строка
                warn_header = doc.add_paragraph("Отклонения от норм (Коррозия и Остаточный ресурс):")
                for run in warn_header.runs:
                    run.font.bold = True
                    run.font.color.rgb = docx.shared.RGBColor(211, 47, 47)
                
                for warn in self.warnings_list:
                    p = doc.add_paragraph(warn)
                    for run in p.runs:
                        run.font.bold = True
                    
            out_path = os.path.normpath(os.path.join(self.local_root, "Итоговый_Отчет_ПНОС.docx"))
            doc.save(out_path)
            
            QMessageBox.information(self, "Успех", f"Отчет успешно сохранен по пути:\n{out_path}")
            
            # Пытаемся открыть файл
            if os.name == 'nt':
                os.startfile(out_path)
                
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить отчет:\n{e}")
